from docx import Document
import uuid

def generate_docx(text: str):
    doc = Document()
    doc.add_heading("AI Doubt Solver Answer", level=1)
    doc.add_paragraph(text)

    path = f"C:/Users/hp/Documents/answer_{uuid.uuid4().hex[:6]}.docx"
    doc.save(path)
    return path
